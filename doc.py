import os
from docx import Document
from langchain.llms import OpenAI
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain_openai import ChatOpenAI
from ai import count_tokens_accurate
from ai import compute_token_cost
from ai import consume_user_tokens

# 创建翻译链
def create_translation_chain(lang):
    llm = OpenAI(temperature=0.7)
    template = """
    将以下文本翻译成{lang}：

    {text}

    翻译：
    """
    prompt = PromptTemplate(input_variables=["lang","text"], template=template)
    return LLMChain(llm=llm, prompt=prompt)

def create_translation_chain_by_dk(lang):
    template = """
    你是一个专业的翻译专家，翻译内容要求词义准确并且简单，字数保持和原来差不多，（数字部分，数学公式，特殊符号,Markdown的符号就不用翻译），不要输出多余内容
    将以下文本翻译成{lang}：
    {text} 
    """
    prompt = PromptTemplate(
        input_variables=["text", "lang"],
        template=template
    )
    ark_key = os.getenv('ARK_API_KEY')
    ark_base = os.getenv('ARK_BASE_URL', 'https://ark.cn-beijing.volces.com/api/v3')
    ark_model = os.getenv('ARK_MODEL')
    if ark_key and ark_model:
        llm = ChatOpenAI(
            openai_api_key=ark_key,
            openai_api_base=ark_base,
            model_name=ark_model,
            temperature=0.3,
            callbacks=[]
        )
    else:
        key = os.getenv('OPENAI_API_KEY')
        if not key or not key.strip():
            raise RuntimeError("OPENAI_API_KEY is not set")
        llm = ChatOpenAI(
            openai_api_key=key,
            model_name="gpt-4o",
            temperature=0.3,
            callbacks=[]
        )
    # 创建翻译链
    chain = LLMChain(llm=llm, prompt=prompt, callbacks=[])
    
    return chain

# 计算总字数
def count_total_words(doc):
    total_words = 0
    for paragraph in doc.paragraphs:
        total_words += len(paragraph.text.split())
    for table in doc.tables:
        total_words += count_words_in_table(table)
    return total_words

# 计算表格中的字数
def count_words_in_table(table):
    total_words = 0
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                total_words += len(paragraph.text.split())
            for nested_table in cell.tables:
                total_words += count_words_in_table(nested_table)
    return total_words

# 翻译段落并保持格式
def translate_paragraph(paragraph, translation_chain,lang):
    if paragraph.text.strip():
        origin = paragraph.text
        translated_text = translation_chain.run(text=paragraph.text,lang=lang).strip()
        runs = paragraph.runs
        paragraph.clear()
        new_run = paragraph.add_run(translated_text)
        if runs:
            original_run = runs[0]
            new_run.font.name = original_run.font.name
            new_run.font.size = original_run.font.size
            new_run.font.bold = original_run.font.bold
            new_run.font.italic = original_run.font.italic
            if original_run.font.color.rgb:
                new_run.font.color.rgb = original_run.font.color.rgb
        input_tokens = count_tokens_accurate(origin.strip(), translation_chain, lang)
        output_tokens = count_tokens_accurate(translated_text, translation_chain, lang)
        token_cost = compute_token_cost(input_tokens, output_tokens)
        return len(origin.split()), token_cost
    return 0, 0

# 处理表格
def process_table(table, translation_chain, progress_callback, task_id, total_words, translated_words,lang, total_tokens_acc):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                translated_count, token_cost = translate_paragraph(paragraph, translation_chain,lang)
                translated_words[0] += translated_count
                total_tokens_acc[0] += token_cost
                progress = translated_words[0] / total_words * 100
                progress_callback(task_id, progress)
            for nested_table in cell.tables:
                process_table(nested_table, translation_chain, progress_callback, task_id, total_words, translated_words,lang, total_tokens_acc)

# 替换文本并保持格式
def replace_text_in_word(task_id, input_docx_path, output_docx_path, progress_callback, lang, user_id=None):
    doc = Document(input_docx_path)
    translation_chain = create_translation_chain_by_dk(lang)

    total_words = count_total_words(doc)
    translated_words = [0]  # 使用列表以便在闭包中更新
    # 用于计算token的变量
    total_tokens = [0]
    for paragraph in doc.paragraphs:
        translated_count, token_cost = translate_paragraph(paragraph, translation_chain,lang)
        translated_words[0] += translated_count
        total_tokens[0] += token_cost
        progress = translated_words[0] / total_words * 100
        progress_callback(task_id, progress)

    for table in doc.tables:
        process_table(table, translation_chain, progress_callback, task_id, total_words, translated_words,lang, total_tokens)

    doc.save(output_docx_path)
    progress_callback(task_id, 100)  # 完成时确保进度设置为100%
    if user_id:
        try:
            consume_user_tokens(user_id, total_tokens[0])
        except Exception:
            pass

def main():
    input_docx_path = 'file/test2.docx'
    output_docx_path = 'output_translated.docx'
    
    def dummy_progress_callback(task_id, progress):
        print(f"Task {task_id} progress: {progress}%")
    
    replace_text_in_word("task1", input_docx_path, output_docx_path, dummy_progress_callback,"English")
    print(f"Word document saved as {output_docx_path} with translated text")

if __name__ == '__main__':
    main()
